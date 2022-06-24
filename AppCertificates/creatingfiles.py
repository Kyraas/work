# -*- coding: utf-8 -*-
from datetime import datetime

from xlsxwriter import Workbook, exceptions
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm
from PyQt6.QtWidgets import QMessageBox

headers = [ '№ п/п', 'Номер сертификата', 'Дата выдачи', 
            'Срок действия', 'Наименование средства (шифр)',
            'Наименования документов, ' + \
            'требованиям которых соответствует средство',
            'Схема сертификации', 'Испытательная лаборатория',
            'Орган по сертификации', 'Заявитель',
            'Реквизиты заявителя (индекс, адрес, телефон)',
            'Информация об окончании срока технической поддержки,' + \
            'полученная от заявителя'
            ]


def save_excel_file(self, fileName, table):

    # Создание excel-файла с указанным названием
    # и форматирование колонок и строки.
    workbook = Workbook(fileName)
    worksheet = workbook.add_worksheet()
    text_wrap = workbook.add_format({'text_wrap':True,
                                    'valign':'vcenter',
                                    'align':'center'})

    head = workbook.add_format({'bold':True,
                                'valign':'vcenter',
                                'align':'center',
                                'text_wrap':True})

    worksheet.set_column(0, 10, None, text_wrap)
    worksheet.set_column(0, 0, 13)
    worksheet.set_column(1, 1, 11)
    worksheet.set_column(2, 2, 16)
    worksheet.set_column(3, 4, 40)
    worksheet.set_column(5, 6, 20)
    worksheet.set_column(7, 7, 30)
    worksheet.set_column(8, 9, 40)
    worksheet.set_column(10, 10, 16)
    worksheet.set_row(0, None, head)

    # Запись заголовков в строку 0
    # Игнорируем первый заголовок, т.к. он для rowid
    for i in range(11):
        worksheet.write(0, i, headers[i+1]) 

    # Запись данных, начиная со строки 1
    for r, row in enumerate(table, start=1):
        for c, cell in enumerate(row):
            try:
                date = datetime.strptime(
                    cell, "%Y-%m-%d").date().strftime("%d.%m.%Y")
                cell = str(date)
            except:
                pass
            # Запись данных в клетку (строка, колонка, данные)
            worksheet.write(r, c, cell)
    try:
        workbook.close()
    except exceptions.FileCreateError:
        QMessageBox.warning(self, "Ошибка доступа",
                            f"Файл '{fileName}' не может быт" +
                            "перезаписан, так как уже открыт.\n" +
                            "Закройте файл и повторите попытку. ")
        return True


def save_word_file(self, fileName, data):

    # добавим заголовки и нумерацию строк в таблицу, начиная от 1
    data.insert(0, headers)
    k = 1
    for i in range(1, len(data)):
        data[i].insert(0, k)
        k += 1

    doc = Document()

    # изменение ориентации страниц (меняем местами ширину и высоту)
    section = doc.sections[-1]  # применяется ко всем листам (-1)
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

    table = doc.add_table(rows = len(data), cols = 12)
    t_cells = table._cells

    # Без этого таблица будет с невидимыми границами
    table.style = 'Table Grid'

    # Добавление остальных данных
    for i in range(len(data)):
        row_cells = t_cells[i*12:(i+1)*12]
        for j in range(12):
            try:
                date = datetime.strptime(
                    data[i][j],"%Y-%m-%d").date().strftime("%d.%m.%Y")

                data[i][j] = str(date)
            except:
                pass
            row_cells[j].text = str(data[i][j])

        # Форматирование таблицы
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                p_format = paragraph.paragraph_format
                p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Calibri'
                    font.size = Pt(7)
                    if i == 0:
                        font.bold = True
    try:
        doc.save(fileName)
    except PermissionError:
        QMessageBox.warning(self, "Ошибка доступа",
                            f"Файл '{fileName}' не может быть" +
                            "перезаписан, так как уже открыт.\n" +
                            "Закройте файл и повторите попытку. ")
        return True
